"""
翻译服务路由

即时翻译、文档翻译、翻译预览与下载
"""
import time
import uuid
import logging
from typing import Optional
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status, Header
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.database import get_db
from app.models.translation import TranslationTask
from app.models.user import User
from app.services.translation_service import translation_service
from app.utils.errors import handle_app_errors, NotFoundException, ValidationException, UnauthorizedException
from docx import Document

logger = logging.getLogger("api")

router = APIRouter(prefix="/api/translate", tags=["翻译服务"])

# 配置
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {".txt", ".docx"}


def generate_task_id() -> str:
    """生成翻译任务ID"""
    return f"trans_{uuid.uuid4().hex[:16]}"


async def get_current_user(
    authorization: Optional[str] = None,
    db: AsyncSession = Depends(get_db)
) -> User:
    """从Authorization头获取当前用户（简化版）"""
    if not authorization or not authorization.startswith("Bearer "):
        raise UnauthorizedException("请先登录")

    token = authorization.replace("Bearer ", "")
    # 简化：这里应该验证JWT token
    # TODO: 接入完整的JWT验证
    result = await db.execute(select(User).where(User.id == int(token)))
    user = result.scalar_one_or_none()

    if not user:
        raise UnauthorizedException("用户不存在")

    return user


@router.post("")
@handle_app_errors
async def translate_text(
    text: str = Form(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    """
    即时翻译

    Request:
        - text: 待翻译文本
        - source_lang: 源语言 (如 en, zh)
        - target_lang: 目标语言 (如 zh, en)

    Response:
        - translated_text: 翻译结果
        - source_lang: 源语言
        - target_lang: 目标语言
    """
    start_time = time.time()
    task_id = generate_task_id()
    logger.info(f"[Translate] Text translation started: task_id={task_id}, text_length={len(text)}, source={source_lang}, target={target_lang}")

    # 文本校验
    if not text or not text.strip():
        raise ValidationException("翻译文本不能为空")

    if len(text) > 5000:
        raise ValidationException("文本长度不能超过5000字符")

    try:
        # 调用翻译服务
        translated_text = await translation_service.translate_text(text, source_lang, target_lang)

        elapsed_ms = int((time.time() - start_time) * 1000)
        logger.info(f"[Translate] Text translation success: task_id={task_id}, elapsed={elapsed_ms}ms, result_length={len(translated_text)}")

        return {
            "success": True,
            "data": {
                "translated_text": translated_text,
                "source_lang": source_lang,
                "target_lang": target_lang
            }
        }

    except Exception as e:
        elapsed_ms = int((time.time() - start_time) * 1000)
        logger.error(f"[Translate] Text translation failed: task_id={task_id}, elapsed={elapsed_ms}ms, error={str(e)}")
        raise


@router.post("/docx")
@handle_app_errors
async def translate_document(
    file: UploadFile = File(...),
    target_lang: str = Form(...),
    source_lang: str = Form("cn"),  # 文档翻译源语言固定为中文
    db: AsyncSession = Depends(get_db),
    authorization: Optional[str] = Header(None)
):
    """
    文档翻译（上传.txt或.docx文件）

    Request (multipart/form-data):
        - file: .txt或.docx文件
        - target_lang: 目标语言 (如 zh, en)
        - source_lang: 源语言，默认为auto

    Response:
        - task_id: 翻译任务ID
        - status: 任务状态
    """
    start_time = time.time()
    task_id = generate_task_id()

    # 文件名脱敏
    safe_filename = file.filename if file.filename else "unknown"
    logger.info(f"[Translate] Document translation started: task_id={task_id}, filename={safe_filename}, target={target_lang}")

    # 文件类型校验
    if not file.filename:
        raise ValidationException("文件名不能为空")

    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise ValidationException(f"不支持的文件类型，仅支持: {', '.join(ALLOWED_EXTENSIONS)}")

    # 文件大小校验
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise ValidationException(f"文件大小不能超过10MB")

    # 读取文件内容
    try:
        if file_ext == ".txt":
            original_content = content.decode("utf-8")
        elif file_ext == ".docx":
            original_content = _read_docx_content(content)
        else:
            raise ValidationException(f"不支持的文件类型: {file_ext}")
    except Exception as e:
        logger.error(f"[Translate] Failed to read file: task_id={task_id}, error={str(e)}")
        raise ValidationException(f"文件读取失败: {str(e)}")

    # 创建翻译任务记录
    task = TranslationTask(
        task_id=task_id,
        user_id=1,  # TODO: 从登录态获取
        task_type="document",
        source_lang=source_lang,
        target_lang=target_lang,
        original_content=original_content,
        status="processing",
        word_count=_count_words(original_content)
    )
    db.add(task)
    await db.commit()
    await db.refresh(task)

    try:
        # 调用翻译服务
        translated_content = await translation_service.translate_document(
            original_content,
            target_lang,
            source_lang
        )

        # 更新任务状态
        task.status = "completed"
        task.translated_content = translated_content
        await db.commit()

        elapsed_ms = int((time.time() - start_time) * 1000)
        logger.info(f"[Translate] Document translation success: task_id={task_id}, elapsed={elapsed_ms}ms, word_count={task.word_count}")

        return {
            "success": True,
            "data": {
                "task_id": task_id,
                "status": "completed"
            }
        }

    except Exception as e:
        # 更新任务失败状态
        task.status = "failed"
        task.error_message = str(e)
        await db.commit()

        elapsed_ms = int((time.time() - start_time) * 1000)
        logger.error(f"[Translate] Document translation failed: task_id={task_id}, elapsed={elapsed_ms}ms, error={str(e)}")
        raise


@router.get("/preview/{task_id}")
@handle_app_errors
async def get_translation_preview(
    task_id: str,
    db: AsyncSession = Depends(get_db),
    authorization: Optional[str] = Header(None)
):
    """
    获取翻译预览

    Response:
        - task_id: 任务ID
        - status: 任务状态 (pending/processing/completed/failed)
        - translated_content: 翻译结果内容
        - word_count: 词数统计
    """
    logger.info(f"[Translate] Get preview: task_id={task_id}")

    result = await db.execute(
        select(TranslationTask).where(TranslationTask.task_id == task_id)
    )
    task = result.scalar_one_or_none()

    if not task:
        raise NotFoundException("TranslationTask", task_id)

    response_data = {
        "task_id": task.task_id,
        "status": task.status,
        "word_count": task.word_count
    }

    if task.status == "completed":
        response_data["translated_content"] = task.translated_content
    elif task.status == "failed":
        response_data["error_message"] = task.error_message

    logger.info(f"[Translate] Preview response: task_id={task_id}, status={task.status}")

    return {
        "success": True,
        "data": response_data
    }


@router.get("/download/{task_id}")
@handle_app_errors
async def download_translated_document(
    task_id: str,
    db: AsyncSession = Depends(get_db),
    authorization: Optional[str] = Header(None)
):
    """
    下载翻译后的Word文档

    Response:
        - .docx文件流
    """
    logger.info(f"[Translate] Download document: task_id={task_id}")

    result = await db.execute(
        select(TranslationTask).where(TranslationTask.task_id == task_id)
    )
    task = result.scalar_one_or_none()

    if not task:
        raise NotFoundException("TranslationTask", task_id)

    if task.status != "completed":
        raise ValidationException(f"翻译任务未完成，当前状态: {task.status}")

    if not task.translated_content:
        raise ValidationException("翻译结果为空")

    # 生成Word文档
    docx_content = _generate_docx(task.translated_content)

    logger.info(f"[Translate] Document generated: task_id={task_id}, size={len(docx_content)} bytes")

    return StreamingResponse(
        iter([docx_content]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="translated_{task_id}.docx"'
        }
    )


def _read_docx_content(content: bytes) -> str:
    """读取.docx文件内容"""
    import io
    doc = Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    return "\n".join(paragraphs)


def _count_words(text: str) -> int:
    """统计词数（简单按空格和中文分词）"""
    import re
    # 去除多余空白
    text = re.sub(r'\s+', ' ', text).strip()
    if not text:
        return 0

    # 统计中文词（每4个字符算一个词）
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    # 统计英文词
    english_words = len(re.findall(r'[a-zA-Z]+', text))

    # 粗略估算：中文每个字算1词，英文每个单词算1词
    return chinese_chars + english_words


def _generate_docx(content: str) -> bytes:
    """生成Word文档"""
    import io
    doc = Document()
    # 添加内容
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    # 保存到bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()