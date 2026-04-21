"""
DOCX文档生成服务

基于python-docx的Word文档生成
"""
import logging
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger("app")


class DocxGeneratorService:
    """DOCX文档生成服务"""

    def __init__(self):
        self.logger = logging.getLogger("app")

    def _is_template_content(self, text: str) -> bool:
        """检测是否为模板占位符内容"""
        import re
        if not text:
            return True
        template_phrases = [
            "详细描述讨论了什么问题",
            "各方观点",
            "最终结论是什么",
            "决定的具体内容",
            "任务内容说明",
            "负责人：XXX",
            "截止时间：XXX",
            "根据转写内容分析",
            "根据转写内容总结",
            "根据转写内容",
        ]
        text_lower = text.lower()
        match_count = sum(1 for phrase in template_phrases if phrase.lower() in text_lower)
        return match_count >= 2

    def _filter_list_field(self, items) -> list:
        """过滤列表字段，移除模板内容"""
        import re
        if not items:
            return []
        if isinstance(items, str):
            items = [items]
        if not isinstance(items, list):
            return [str(items)]

        result = []
        for item in items:
            item_str = str(item).strip()
            if not item_str or len(item_str) < 10:
                continue
            # 清理代码格式
            item_str = re.sub(r"\{[^}]*\}", lambda m: m.group(0).replace("{", "").replace("}", "").replace("'", ""), item_str)
            item_str = item_str.replace("{", "").replace("}", "").replace("'", "").strip()
            # 过滤模板内容
            if item_str and not self._is_template_content(item_str) and len(item_str) > 10:
                result.append(item_str)
        return result

    def _add_title(self, doc: Document, text: str, level: int = 1) -> None:
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_section_heading(self, doc: Document, text: str) -> None:
        """添加章节标题"""
        heading = doc.add_heading(text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False) -> None:
        """添加段落"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_bullet_point(self, doc: Document, text: str, level: int = 0) -> None:
        """添加项目符号"""
        para = doc.add_paragraph(text, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        for run in para.runs:
            run.font.size = Pt(11)

    def _format_bytes_size(self, size: int) -> str:
        """格式化字节大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.2f} {unit}"
            size /= 1024.0
        return f"{size:.2f} TB"

    async def generate_course_docx(self, summary: Dict[str, Any]) -> bytes:
        """
        生成课程总结DOCX

        包含：主题、核心知识点、重点难点、金句、预习建议

        Args:
            summary: 课程总结数据字典
                - topic: 主题
                - key_points: 核心知识点 (str or list)
                - difficulties: 重点难点 (str or list)
                - memorable_quote: 金句
                - next_suggestion: 预习建议
                - full_text: 完整总结 (可选)
                - course_name: 课程名称 (可选)
                - teacher: 教师 (可选)

        Returns:
            DOCX文件字节数据
        """
        import re

        start_time = datetime.now()
        topic = summary.get("topic", "未命名课程")

        # 预处理：清理代码格式
        def clean_text(text):
            if not text:
                return text
            text = re.sub(r"\{[^}]*\}", lambda m: m.group(0).replace("{", "").replace("}", "").replace("'", ""), str(text))
            text = text.replace("{", "").replace("}", "").replace("'", "").strip()
            return text

        for key in ['topic', 'key_points', 'difficulties', 'memorable_quote', 'next_suggestion', 'full_text']:
            if key in summary and summary[key]:
                summary[key] = clean_text(summary[key])

        self.logger.info(f"[DOCX] 开始生成课程总结文档, 主题: {topic}")
        self.logger.debug(f"[DOCX] 课程总结数据: {summary}")

        try:
            # 创建文档
            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/main}eastAsia', '微软雅黑')
            style.font.size = Pt(11)

            # 添加主标题
            self._add_title(doc, "课程总结", level=1)

            # 添加副标题（主题）
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(f"主题：{topic}")
            subtitle_run.bold = True
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(0x08, 0x91, 0xB2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加课程信息（如果存在）
            if summary.get("course_name") or summary.get("teacher"):
                info_parts = []
                if summary.get("course_name"):
                    info_parts.append(f"课程：{summary['course_name']}")
                if summary.get("teacher"):
                    info_parts.append(f"教师：{summary['teacher']}")
                info_para = doc.add_paragraph()
                info_run = info_para.add_run(" | ".join(info_parts))
                info_run.font.size = Pt(10)
                info_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加时间戳
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # 空行

            # 核心知识点
            self._add_section_heading(doc, "一、核心知识点")
            key_points = summary.get("key_points", "")
            if isinstance(key_points, list):
                for point in key_points:
                    self._add_bullet_point(doc, str(point))
            elif isinstance(key_points, str) and key_points:
                # 按换行符分割字符串格式的列表
                points = [p.strip() for p in key_points.split('\n') if p.strip()]
                if points:
                    for point in points:
                        self._add_bullet_point(doc, point)
                else:
                    self._add_paragraph(doc, str(key_points))
            elif key_points:
                self._add_paragraph(doc, str(key_points))
            else:
                self._add_paragraph(doc, "暂无")

            # 重点难点
            self._add_section_heading(doc, "二、重点难点")
            difficulties = summary.get("difficulties", "")
            if isinstance(difficulties, list):
                for point in difficulties:
                    self._add_bullet_point(doc, str(point))
            elif isinstance(difficulties, str) and difficulties:
                # 按换行符分割字符串格式的列表
                points = [p.strip() for p in difficulties.split('\n') if p.strip()]
                if points:
                    for point in points:
                        self._add_bullet_point(doc, point)
                else:
                    self._add_paragraph(doc, str(difficulties))
            elif difficulties:
                self._add_paragraph(doc, str(difficulties))
            else:
                self._add_paragraph(doc, "暂无")

            # 金句
            self._add_section_heading(doc, "三、金句")
            memorable_quote = summary.get("memorable_quote", "")
            if memorable_quote:
                quote_para = doc.add_paragraph()
                quote_run = quote_para.add_run(f'"{memorable_quote}"')
                quote_run.italic = True
                quote_run.font.size = Pt(12)
                quote_run.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
                quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加引用框背景效果（通过边框模拟）
                quote_para.paragraph_format.space_before = Pt(12)
                quote_para.paragraph_format.space_after = Pt(12)
            else:
                self._add_paragraph(doc, "暂无")

            # 预习建议
            self._add_section_heading(doc, "四、预习建议")
            next_suggestion = summary.get("next_suggestion", "")
            if next_suggestion:
                # 添加带边框的建议框
                self._add_paragraph(doc, str(next_suggestion))
            else:
                self._add_paragraph(doc, "暂无")

            # 完整总结（如果存在）
            if summary.get("full_text"):
                doc.add_paragraph()
                self._add_section_heading(doc, "五、完整总结")
                full_text = summary.get("full_text", "")
                if isinstance(full_text, list):
                    for paragraph in full_text:
                        para = doc.add_paragraph()
                        run = para.add_run(str(paragraph))
                        run.font.size = Pt(11)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.space_after = Pt(8)
                else:
                    # 分段落显示完整总结
                    paragraphs = str(full_text).split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph()
                            run = para.add_run(para_text.strip())
                            run.font.size = Pt(11)
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            para.paragraph_format.space_after = Pt(8)

            # 保存到字节流
            byte_stream = BytesIO()
            doc.save(byte_stream)
            byte_stream.seek(0)
            content = byte_stream.getvalue()

            # 记录日志
            duration = (datetime.now() - start_time).total_seconds()
            file_size = self._format_bytes_size(len(content))

            self.logger.info(
                f"[DOCX] 课程总结文档生成完成, "
                f"主题: {topic}, "
                f"文件大小: {file_size}, "
                f"耗时: {duration:.3f}秒"
            )

            return content

        except Exception as e:
            self.logger.error(
                f"[DOCX] 课程总结文档生成失败, 主题: {topic}, 错误: {str(e)}",
                exc_info=True
            )
            raise

    async def generate_meeting_docx(self, summary: Dict[str, Any]) -> bytes:
        """
        生成会议总结DOCX

        包含：主题、讨论要点、决议事项、待办事项

        Args:
            summary: 会议总结数据字典
                - topic: 主题
                - discussion_points: 讨论要点 (str or list)
                - decisions: 决议事项 (str or list)
                - action_items: 待办事项 (str or list)
                - meeting_date: 会议日期 (可选)
                - attendees: 与会人员 (可选)
                - full_text: 完整总结 (可选)

        Returns:
            DOCX文件字节数据
        """
        import re

        start_time = datetime.now()
        topic = summary.get("topic", "未命名会议")

        # 预处理：清理代码格式
        def clean_text(text):
            if not text:
                return text
            text = re.sub(r"\{[^}]*\}", lambda m: m.group(0).replace("{", "").replace("}", "").replace("'", ""), str(text))
            text = text.replace("{", "").replace("}", "").replace("'", "").strip()
            return text

        for key in ['topic', 'discussion_points', 'resolutions', 'decisions', 'action_items', 'memorable_quote', 'full_text']:
            if key in summary and summary[key]:
                summary[key] = clean_text(summary[key])

        self.logger.info(f"[DOCX] 开始生成会议总结文档, 主题: {topic}")
        self.logger.debug(f"[DOCX] 会议总结数据: {summary}")

        try:
            # 创建文档
            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/main}eastAsia', '微软雅黑')
            style.font.size = Pt(11)

            # 添加主标题
            self._add_title(doc, "会议总结", level=1)

            # 添加副标题（主题）
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(f"主题：{topic}")
            subtitle_run.bold = True
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(0x08, 0x91, 0xB2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加会议信息（如果存在）
            info_parts = []
            if summary.get("meeting_date"):
                info_parts.append(f"日期：{summary['meeting_date']}")
            if summary.get("attendees"):
                attendees = summary["attendees"]
                if isinstance(attendees, list):
                    attendees = "、".join(attendees)
                info_parts.append(f"与会人员：{attendees}")

            if info_parts:
                info_para = doc.add_paragraph()
                info_run = info_para.add_run(" | ".join(info_parts))
                info_run.font.size = Pt(10)
                info_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加时间戳
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # 空行

            # 讨论要点
            self._add_section_heading(doc, "一、讨论要点")
            discussion_points = self._filter_list_field(summary.get("discussion_points", ""))
            if discussion_points:
                for i, point in enumerate(discussion_points, 1):
                    self._add_bullet_point(doc, f"{i}. {point}")
            else:
                self._add_paragraph(doc, "暂无")

            # 决议事项
            self._add_section_heading(doc, "二、决议事项")
            resolutions = self._filter_list_field(summary.get("resolutions", ""))
            if resolutions:
                for i, item in enumerate(resolutions, 1):
                    self._add_bullet_point(doc, f"{i}. {item}")
            else:
                self._add_paragraph(doc, "暂无")

            # 待办事项
            self._add_section_heading(doc, "三、待办事项")
            action_items = self._filter_list_field(summary.get("action_items", ""))
            if action_items:
                for i, item in enumerate(action_items, 1):
                    self._add_bullet_point(doc, f"{i}. {item}")
            else:
                self._add_paragraph(doc, "暂无")

            # 完整总结（如果存在）
            if summary.get("full_text"):
                doc.add_paragraph()
                self._add_section_heading(doc, "四、完整记录")
                full_text = summary.get("full_text", "")
                if isinstance(full_text, list):
                    for paragraph in full_text:
                        self._add_paragraph(doc, str(paragraph))
                else:
                    self._add_paragraph(doc, str(full_text))

            # 保存到字节流
            byte_stream = BytesIO()
            doc.save(byte_stream)
            byte_stream.seek(0)
            content = byte_stream.getvalue()

            # 记录日志
            duration = (datetime.now() - start_time).total_seconds()
            file_size = self._format_bytes_size(len(content))

            self.logger.info(
                f"[DOCX] 会议总结文档生成完成, "
                f"主题: {topic}, "
                f"文件大小: {file_size}, "
                f"耗时: {duration:.3f}秒"
            )

            return content

        except Exception as e:
            self.logger.error(
                f"[DOCX] 会议总结文档生成失败, 主题: {topic}, 错误: {str(e)}",
                exc_info=True
            )
            raise


    async def generate_review_summary(
        self,
        title: str,
        topic: str,
        key_points: str,
        difficulties: str,
        memorable_quote: str,
        next_suggestion: str,
        full_text: str,
        metadata: dict = None
    ) -> bytes:
        """
        生成录音回顾总结文档

        Args:
            title: 标题
            topic: 主题
            key_points: 核心知识点（课程用）或讨论要点（会议用）
            difficulties: 重点难点（课程用）或决议事项（会议用）
            memorable_quote: 金句/重要发言
            next_suggestion: 预习建议（课程用）或待办事项（会议用）
            full_text: 完整总结
            metadata: 额外元数据 (record_type: course/meeting)

        Returns:
            DOCX文件字节数据
        """
        record_type = metadata.get("record_type", "course") if metadata else "course"

        if record_type == "meeting":
            # 会议类型使用会议专用字段
            summary = {
                "topic": topic,
                "discussion_points": key_points,  # 讨论要点
                "resolutions": difficulties,     # 决议事项
                "action_items": next_suggestion,  # 待办事项
                "memorable_quote": memorable_quote,
                "full_text": full_text,
                "meeting_date": metadata.get("created_at", "") if metadata else ""
            }
            return await self.generate_meeting_docx(summary)
        else:
            # 课程类型
            summary = {
                "topic": topic,
                "key_points": key_points,
                "difficulties": difficulties,
                "memorable_quote": memorable_quote,
                "next_suggestion": next_suggestion,
                "full_text": full_text,
                "course_name": title,
            }
            return await self.generate_course_docx(summary)


# 单例实例
docx_generator = DocxGeneratorService()